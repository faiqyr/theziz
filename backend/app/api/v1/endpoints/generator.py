import io
from typing import List, Optional

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from fastapi import APIRouter
from fastapi.responses import StreamingResponse
from pydantic import BaseModel

router = APIRouter()

# --- SCHEMA DATA BLOK KONTEN BARU ---
class KontenBlok(BaseModel):
    tipe: str                 # "heading" atau "paragraf"
    level_heading: Optional[int] = 2  # 1, 2, 3, atau 4
    nomor: Optional[str] = ""         # Contoh: "2.1", "2.1.1", "a"
    teks: str

class DocConfig(BaseModel):
    jenis_font: str
    ukuran_font: int
    jarak_spasi: float
    margin_top: float      
    margin_bottom: float   
    margin_left: float     
    margin_right: float    
    pakai_cover: bool  
    pakai_kata_pengantar: bool  
    pakai_lampiran: bool         
    
    number_format: str           
    include_chapter_number: bool 
    nomor_bab_aktif: str         
    separator_type: str          
    
    isi_teks: str
    
    # Menambahkan penampung array blok dari antarmuka frontend (default kosong)
    blok_naskah: List[KontenBlok] = []

# --- ENGINE: CUSTOM PAGE NUMBER BUILDER ---
def add_custom_page_number(paragraph, font_name, prefix_bab="", separator="-", format_style="angka", space_around_sep=True):
    if prefix_bab:
        run_prefix = paragraph.add_run(f"{prefix_bab}")
        run_prefix.font.name = font_name
        run_prefix.font.size = Pt(11)
        
        if separator != "none":
            sep_text = f" {separator} " if space_around_sep else separator
            run_sep = paragraph.add_run(sep_text)
            run_sep.font.name = font_name
            run_sep.font.size = Pt(11)

    if format_style == "strip" and not prefix_bab:
        run_open = paragraph.add_run("-")
        run_open.font.name = font_name
        run_open.font.size = Pt(11)

    run_field = paragraph.add_run()
    run_field.font.name = font_name
    run_field.font.size = Pt(11)
    
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    
    if format_style == "huruf":
        instrText.text = "PAGE \\* alphabetic" 
    elif format_style == "romawi":
        instrText.text = "PAGE \\* roman" 
    else:
        instrText.text = "PAGE" 
        
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    
    run_field._r.append(fldChar1)
    run_field._r.append(instrText)
    run_field._r.append(fldChar2)
    run_field._r.append(fldChar3)

    if format_style == "strip" and not prefix_bab:
        run_close = paragraph.add_run("-")
        run_close.font.name = font_name
        run_close.font.size = Pt(11)

def add_page_number_romawi_awal(run):
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE \\* roman"  
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)

def set_section_page_number_format(section, fmt_type="decimal", start_from=None):
    sectPr = section._sectPr
    pgNumType = sectPr.find(qn('w:pgNumType'))
    if pgNumType is None:
        pgNumType = OxmlElement('w:pgNumType')
        sectPr.append(pgNumType)
    pgNumType.set(qn('w:fmt'), fmt_type) 
    if start_from is not None:
        pgNumType.set(qn('w:start'), str(start_from))

@router.post("/generate-font-doc")
async def generate_font_doc(config: DocConfig):
    doc = Document()

    # --- KONFIGURASI HEADING STYLE (Fokus Utama Fitur Baru) ---
    for i in range(1, 5):
        h_style = doc.styles[f'Heading {i}']
        h_style.font.name = config.jenis_font
        h_style.font.color.rgb = None  # Warna hitam otomatis
        h_style.font.bold = True
        h_style.paragraph_format.first_line_indent = Cm(0)
        h_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        
        if i == 1:
            h_style.font.size = Pt(16)
            h_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            h_style.font.size = Pt(12)
            h_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
    # Global Style Default (Applies to Chapters and Appendices)
    style_normal = doc.styles['Normal']
    style_normal.font.name = config.jenis_font
    style_normal.font.size = Pt(config.ukuran_font)
    if config.jarak_spasi == 1.0: style_normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    elif config.jarak_spasi == 1.5: style_normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    elif config.jarak_spasi == 2.0: style_normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    else: style_normal.paragraph_format.line_spacing = Pt(config.jarak_spasi * config.ukuran_font)
    
    style_normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    style_normal.paragraph_format.first_line_indent = Cm(1.27)

    # ==========================================
    # SECTION 1: FRONT MATTER (COVER & PREFACE)
    # ==========================================
    section_awal = doc.sections[0]
    section_awal.top_margin = Cm(config.margin_top)
    section_awal.bottom_margin = Cm(config.margin_bottom)
    section_awal.left_margin = Cm(config.margin_left)
    section_awal.right_margin = Cm(config.margin_right)
    section_awal.different_first_page_header_footer = True
    
    set_section_page_number_format(section_awal, fmt_type="lrLetters", start_from=1)
    
    footer_normal_awal = section_awal.footer.paragraphs[0] 
    footer_normal_awal.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_num_awal = footer_normal_awal.add_run()
    run_num_awal.font.name = config.jenis_font
    run_num_awal.font.size = Pt(11)
    add_page_number_romawi_awal(run_num_awal)
    
    if config.pakai_cover:
        style_cover = doc.add_paragraph()
        style_cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
        style_cover.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        style_cover.paragraph_format.first_line_indent = Cm(0)
        
        run_judul = style_cover.add_run(
            "PERBANDINGAN JSON-STRUCTURED PROMPTING\n"
            "DENGAN UNSTRUCTURED PROMPTING UNTUK\n"
            "MENGHASILKAN VIDEO PADA PRODUKSI ANIMASI\n"
            "PENDEK BERBASIS AI GENERATIF\n\n\n"
        )
        run_judul.font.name = config.jenis_font; run_judul.font.size = Pt(14); run_judul.bold = True
        
        run_sub = style_cover.add_run("USULAN PENELITIAN\nTUGAS AKHIR\n\n\n\n")
        run_sub.font.name = config.jenis_font; run_sub.font.size = Pt(12); run_sub.bold = True
        
        run_mhs = style_cover.add_run("Oleh :\nNama : Muhammad Faiq Yusuf Raharja\nNPM : 227006127\n\n\n\n\n")
        run_mhs.font.name = config.jenis_font; run_mhs.font.size = Pt(12); run_mhs.bold = True
        
        run_logo = style_cover.add_run("[LOGO UNIVERSITAS SILIWANGI]\n\n\n\n\n")
        run_logo.font.name = config.jenis_font; run_logo.font.size = Pt(12); run_logo.bold = True
        
        run_kampus = style_cover.add_run("JURUSAN INFORMATIKA\nFAKULTAS TEKNIK UNIVERSITAS SILIWANGI\nTASIKMALAYA\n2026")
        run_kampus.font.name = config.jenis_font; run_kampus.font.size = Pt(12); run_kampus.bold = True
        
        if config.pakai_kata_pengantar:
            doc.add_page_break()

    if config.pakai_kata_pengantar:
        kp_title_p = doc.add_paragraph()
        kp_title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        kp_title_p.paragraph_format.first_line_indent = Cm(0)
        kp_title_run = kp_title_p.add_run("KATA PENGANTAR\n\n")
        kp_title_run.bold = True
        
        teks_baku_kp = (
            "Puji dan syukur penulis panjatkan ke hadirat Tuhan Yang Maha Esa atas segala rahmat dan karunia-Nya "
            "sehingga penyusunan Usulan Penelitian Tugas Akhir dapat diselesaikan dengan baik.\n\n"
            "Kota Jonggol, Januari 2026\n\n"
            "Penulis"
        )
        for p_text in teks_baku_kp.split('\n\n'):
            if p_text.strip():
                p = doc.add_paragraph(p_text.strip())
                if "Kota" in p_text or "Penulis" in p_text:
                    p.paragraph_format.first_line_indent = Cm(0)
                    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # ==========================================
    # SECTION 2 TO 6: CHAPTERS I TO V LOOP
    # ==========================================
    DAFTAR_BAB = [
        ("I", "PENDAHULUAN"), ("II", "TINJAUAN PUSTAKA"), 
        ("III", "METODOLOGI PENELITIAN"), ("IV", "HASIL DAN PEMBAHASAN"), 
        ("V", "KESIMPULAN DAN SARAN")
    ]

    for index, (romawi_bab, judul_bab) in enumerate(DAFTAR_BAB):
        section_isi = doc.add_section(WD_SECTION_START.NEW_PAGE)
        section_isi.header.is_linked_to_previous = False
        section_isi.footer.is_linked_to_previous = False
        section_isi.first_page_header.is_linked_to_previous = False
        section_isi.first_page_footer.is_linked_to_previous = False
        
        section_isi.top_margin = Cm(config.margin_top)
        section_isi.bottom_margin = Cm(config.margin_bottom)
        section_isi.left_margin = Cm(config.margin_left)
        section_isi.right_margin = Cm(config.margin_right)
        section_isi.different_first_page_header_footer = True
        
        set_section_page_number_format(section_isi, fmt_type="decimal", start_from=1)
        
        prefix_aktif = romawi_bab if config.include_chapter_number else ""
        
        f_p = section_isi.first_page_footer.paragraphs[0]
        f_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_custom_page_number(f_p, config.jenis_font, prefix_aktif, config.separator_type, config.number_format, space_around_sep=True)
        
        h_p = section_isi.header.paragraphs[0]
        h_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        add_custom_page_number(h_p, config.jenis_font, prefix_aktif, config.separator_type, config.number_format, space_around_sep=True)
        
        bab_title_p = doc.add_paragraph()
        bab_title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        bab_title_p.paragraph_format.first_line_indent = Cm(0)
        
        b_run = bab_title_p.add_run(f"BAB {romawi_bab}\n{judul_bab}\n\n")
        b_run.bold = True
        
        # --- RENDER BLOK DINAMIS (JIKA ADA DARI FRONTEND) ---
        if config.blok_naskah:
            for blok in config.blok_naskah:
                if blok.tipe == "heading":
                    if blok.level_heading == 1:
                        doc.add_heading(f"{blok.nomor} {blok.teks.upper()}", level=1)
                    else:
                        doc.add_heading(f"{blok.nomor} {blok.teks}", level=blok.level_heading)
                elif blok.tipe == "paragraf":
                    doc.add_paragraph(blok.teks)
        else:
            # Fallback jika blok kosong
            doc.add_paragraph(config.isi_teks)
            doc.add_paragraph(config.isi_teks)

    # ==========================================
    # SECTION 7: APPENDICES (LAMPIRAN) LOOP
    # ==========================================
    if config.pakai_lampiran:
        DAFTAR_LAMPIRAN = [
            (1, "Dokumentasi Struktur Prompt JSON"),
            (2, "Hasil Evaluasi Metrik LPIPS Video")
        ]
        
        for lamp_num, judul_lamp in DAFTAR_LAMPIRAN:
            section_lamp = doc.add_section(WD_SECTION_START.NEW_PAGE)
            section_lamp.header.is_linked_to_previous = False
            section_lamp.footer.is_linked_to_previous = False
            
            section_lamp.different_first_page_header_footer = False
            
            section_lamp.top_margin = Cm(config.margin_top)
            section_lamp.bottom_margin = Cm(config.margin_bottom)
            section_lamp.left_margin = Cm(config.margin_left)
            section_lamp.right_margin = Cm(config.margin_right)
            
            set_section_page_number_format(section_lamp, fmt_type="decimal", start_from=1)
            
            h_p_lamp = section_lamp.header.paragraphs[0]
            h_p_lamp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            add_custom_page_number(
                h_p_lamp, 
                config.jenis_font, 
                prefix_bab=f"L{lamp_num}", 
                separator="-", 
                format_style="angka", 
                space_around_sep=False
            )
            
            lamp_title_p = doc.add_paragraph()
            lamp_title_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            lamp_title_p.paragraph_format.first_line_indent = Cm(0)
            
            l_run = lamp_title_p.add_run(f"Lampiran {lamp_num}\n{judul_lamp}\n\n")
            l_run.bold = True
            
            doc.add_paragraph(config.isi_teks)
            doc.add_paragraph(config.isi_teks) 

    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    
    return StreamingResponse(
        file_stream, 
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", 
        headers={'Content-Disposition': 'attachment; filename="Format_Skripsi_Custom_Bab1sd5_Lampiran.docx"'}
    )
